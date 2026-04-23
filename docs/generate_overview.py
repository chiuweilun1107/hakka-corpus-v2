#!/usr/bin/env python3
"""產生《臺灣客語語料庫》項目總覽表（附1+附2合計）docx — 橫式 A4"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
ITEMS_FILE = BASE_DIR / "items.json"
OUTPUT_FILE = BASE_DIR / "項目總覽_附1附2合計.docx"

FONT_CJK   = "標楷體"
FONT_LATIN = "Times New Roman"
FONT_SIZE  = Pt(11)
FONT_TITLE = Pt(14)

# 類別顯示順序
CATEGORY_ORDER = [
    "語料檢索系統", "語料視覺化", "語料編輯系統", "語料內容擴充", "例句庫功能",
    "首頁功能", "影音展示",
    "AI 輸入辨識", "AI 語言分析", "AI 整合應用", "AI 核心研發", "語音 AI",
    "教育應用", "教育輔助",
    "互動沉浸體驗",
    "社群推廣", "文創推廣", "社會服務",
    "數位典藏", "平台基礎設施",
]

# 類別區塊底色（交替淺藍 / 淺橘，hex → (R,G,B)）
BLOCK_COLORS = [
    (207, 226, 255),  # 淺藍
    (255, 230, 207),  # 淺橘
    (207, 255, 215),  # 淺綠
    (255, 207, 230),  # 淺粉
    (230, 207, 255),  # 淺紫
    (255, 255, 207),  # 淺黃
    (207, 248, 255),  # 淺青
]

DIFFICULTY_ORDER = {"易": 0, "中": 1, "難": 2}


def apply_font(run, size=FONT_SIZE, bold=False, color=None):
    run.font.name  = FONT_LATIN
    run.font.size  = size
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),   FONT_LATIN)
    rFonts.set(qn("w:hAnsi"),   FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.insert(0, rFonts)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "888888")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"),    str(int(w * 567)))
            tcW.set(qn("w:type"), "dxa")


def write_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT,
               bold=False, size=FONT_SIZE, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    apply_font(run, size=size, bold=bold, color=color)


def set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    w, h = section.page_width, section.page_height
    if w < h:
        section.page_width, section.page_height = h, w
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)


def format_amount(amount):
    return f"NT$ {amount:,.0f}"


def main():
    data    = json.loads(ITEMS_FILE.read_text(encoding="utf-8"))
    pricing = data["pricing"]
    items   = data["items"]

    # 依類別分組，類別內按難度→月數排序
    by_cat = {}
    for item in items:
        by_cat.setdefault(item["category"], []).append(item)
    for cat in by_cat:
        by_cat[cat].sort(key=lambda x: (DIFFICULTY_ORDER[x["difficulty"]], x["months"]))

    # 類別順序（CATEGORY_ORDER 未列者排後面）
    remaining = [c for c in by_cat if c not in CATEGORY_ORDER]
    ordered_cats = [c for c in CATEGORY_ORDER if c in by_cat] + remaining

    # ── 建立 doc ─────────────────────────────────────────────
    doc = Document()
    set_landscape(doc)

    # 主標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《臺灣客語語料庫》未來發展規劃 — 項目總覽（附檔1＋附檔2）")
    apply_font(r, size=FONT_TITLE, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"合計 {len(items)} 項，全案估價 NT$ {sum(int(x['months']*pricing[x['difficulty']]) for x in items):,.0f} 元")
    apply_font(r2, size=Pt(11))

    doc.add_paragraph()

    # ── 欄位定義 ────────────────────────────────────────────
    # 橫式 A4 可用寬 ≈ 29.7 - 5.08 = 24.62 cm
    COL_HEADERS = ["項次", "類別", "項目名稱", "來源", "難度", "月數", "估價（元）", "已驗證"]
    COL_WIDTHS  = [ 1.2,   3.5,    7.0,        3.5,    1.2,   1.2,   3.5,          2.0]
    # 對齊
    COL_ALIGN = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]

    # 計算總列數：表頭 + 每類別 (類頭列 + items + 小計列)
    total_rows = 1
    for cat in ordered_cats:
        total_rows += 1 + len(by_cat[cat]) + 1

    table = doc.add_table(rows=total_rows, cols=len(COL_HEADERS))
    set_table_borders(table)
    set_col_widths(table, COL_WIDTHS)

    row_idx = 0
    HEADER_BG = (30, 60, 114)    # 深藍
    SUBTOT_BG = (240, 240, 240)  # 淺灰

    # 表頭列
    for j, hdr in enumerate(COL_HEADERS):
        cell = table.rows[row_idx].cells[j]
        set_cell_bg(cell, HEADER_BG)
        write_cell(cell, hdr, align=WD_ALIGN_PARAGRAPH.CENTER,
                   bold=True, size=FONT_SIZE, color=(255, 255, 255))
    row_idx += 1

    seq = 0
    for cat_idx, cat in enumerate(ordered_cats):
        cat_items = by_cat[cat]
        block_color = BLOCK_COLORS[cat_idx % len(BLOCK_COLORS)]
        cat_total   = sum(int(x["months"] * pricing[x["difficulty"]]) for x in cat_items)

        # 類別標題列（底色，合併後效果靠底色呈現）
        cat_row = table.rows[row_idx]
        for j in range(len(COL_HEADERS)):
            set_cell_bg(cat_row.cells[j], block_color)
        write_cell(cat_row.cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        write_cell(cat_row.cells[1], f"【{cat}】", bold=True)
        write_cell(cat_row.cells[2], f"{len(cat_items)} 項", bold=False)
        write_cell(cat_row.cells[6], format_amount(cat_total),
                   align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
        row_idx += 1

        # 各項目列
        for item in cat_items:
            seq += 1
            amount = int(item["months"] * pricing[item["difficulty"]])
            src_str = "、".join(item.get("source", []))
            months  = item["months"]
            months_str = f"{months:.0f}" if months == int(months) else str(months)

            data_row = table.rows[row_idx]
            vals = [
                str(seq),
                item["category"],
                item["name"],
                src_str,
                item["difficulty"],
                months_str,
                format_amount(amount),
                "✓" if item.get("verified") else "—",
            ]
            for j, (val, align) in enumerate(zip(vals, COL_ALIGN)):
                write_cell(data_row.cells[j], val, align=align)
            row_idx += 1

        # 小計列
        sub_row = table.rows[row_idx]
        for j in range(len(COL_HEADERS)):
            set_cell_bg(sub_row.cells[j], SUBTOT_BG)
        write_cell(sub_row.cells[2], f"小計：{len(cat_items)} 項",
                   align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
        write_cell(sub_row.cells[6], format_amount(cat_total),
                   align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
        row_idx += 1

    doc.save(str(OUTPUT_FILE))
    print(f"✓ 已產出：{OUTPUT_FILE}")
    print(f"  類別數：{len(ordered_cats)}")
    print(f"  項目數：{len(items)}")
    grand = sum(int(x['months']*pricing[x['difficulty']]) for x in items)
    print(f"  全案合計：NT$ {grand:,.0f} 元")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""產生《臺灣客語語料庫》未來發展規劃-哈瑪星提供 docx"""

import json, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = Path(__file__).parent
IMAGES_DIR = BASE_DIR / "images"
ITEMS_FILE = BASE_DIR / "items.json"
OUTPUT_FILE = BASE_DIR / "哈瑪星回覆_20260423.docx"

# ── 字型常數 ────────────────────────────────────────────────
FONT_CJK = "標楷體"
FONT_LATIN = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_TITLE = Pt(14)

DIFFICULTY_ORDER = {"易": 0, "中": 1, "難": 2}
CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
           "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八",
           "十九", "二十", "二十一", "二十二", "二十三", "二十四", "二十五",
           "二十六", "二十七", "二十八", "二十九", "三十", "三十一", "三十二",
           "三十三", "三十四", "三十五", "三十六", "三十七", "三十八", "三十九",
           "四十", "四十一", "四十二", "四十三", "四十四", "四十五", "四十六",
           "四十七", "四十八", "四十九", "五十", "五十一", "五十二", "五十三",
           "五十四", "五十五", "五十六", "五十七", "五十八", "五十九", "六十",
           "六十一", "六十二", "六十三"]


def set_font(run, size=FONT_SIZE_BODY, bold=False):
    run.font.name = FONT_LATIN
    run.font.size = size
    run.font.bold = bold
    run._r.get_or_add_rPr().append(
        OxmlElement("w:rFonts") if False else _make_rFonts(FONT_CJK)
    )


def _make_rFonts(cjk_font):
    el = OxmlElement("w:rFonts")
    el.set(qn("w:eastAsia"), cjk_font)
    return el


def apply_font(run, size=FONT_SIZE_BODY, bold=False):
    """套用標楷體（CJK）+ Times New Roman（Latin）"""
    run.font.name = FONT_LATIN
    run.font.size = size
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    # 移除舊 rFonts，重建
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.insert(0, rFonts)


def add_para(doc, text="", bold=False, size=FONT_SIZE_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        apply_font(run, size=size, bold=bold)
    return p


def set_table_borders(table):
    """設定表格全框線 single 0.75pt"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_col_widths(table, widths_cm):
    for i, row in enumerate(table.rows):
        for j, (cell, w) in enumerate(zip(row.cells, widths_cm)):
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            # 1 cm = 567 twips
            twips = int(w * 567)
            tcW.set(qn("w:w"), str(twips))
            tcW.set(qn("w:type"), "dxa")


def add_header(doc, text):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        for r in p.runs:
            r.text = ""
    if not header.paragraphs:
        hp = header.add_paragraph()
    else:
        hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(text)
    apply_font(run, size=Pt(12), bold=False)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    apply_font(run, size=Pt(12))
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def format_amount(amount):
    return f"新臺幣 {amount:,.0f} 元"


def main():
    data = json.loads(ITEMS_FILE.read_text(encoding="utf-8"))
    pricing = data["pricing"]  # {"易": 100000, "中": 130000, "難": 170000}
    items = data["items"]

    # 有把握（verified）跨難度全優先 → 難度易→中→難 → months 升冪
    items.sort(key=lambda x: (0 if x.get("verified") else 1, DIFFICULTY_ORDER[x["difficulty"]], x["months"]))

    doc = Document()
    set_page_margins(doc)
    add_page_number_footer(doc)

    # ── 主標題 ─────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("《臺灣客語語料庫》未來發展規劃-哈瑪星提供")
    apply_font(r, size=FONT_SIZE_TITLE, bold=True)

    add_para(doc)  # 空行

    # ── 各項目段落 ──────────────────────────────────────────
    five_year_items = []

    for idx, item in enumerate(items):
        num = CN_NUMS[idx] if idx < len(CN_NUMS) else str(idx + 1)
        name = item["name"]
        difficulty = item["difficulty"]
        months = item["months"]
        verified = item["verified"]
        images = item.get("images", [])
        desc_func = item.get("description_func", "")
        desc_poc = item.get("description_poc", "")
        unit_price = pricing[difficulty]
        total = int(months * unit_price)

        # 五年內可完成（≤60 個月者納入）
        if months <= 60:
            five_year_items.append({
                "seq": idx + 1,
                "name": name,
                "months": months,
                "total": total,
            })

        # ── 項目標題：「一、[項目名稱]」 ──────────────────────
        p_head = doc.add_paragraph()
        if idx > 0:
            p_head.paragraph_format.page_break_before = True
        p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_num = p_head.add_run(f"{num}、")
        apply_font(r_num, size=FONT_SIZE_BODY, bold=False)
        r_name = p_head.add_run(name)
        apply_font(r_name, size=FONT_SIZE_BODY, bold=True)

        # ── 功能說明 ────────────────────────────────────────
        p_func = doc.add_paragraph()
        r_label = p_func.add_run("功能說明：")
        apply_font(r_label, size=FONT_SIZE_BODY)
        r_content = p_func.add_run(desc_func)
        apply_font(r_content, size=FONT_SIZE_BODY)

        # ── 驗證成果（僅已驗證項目輸出） ───────────────────────
        if verified and desc_poc:
            p_poc = doc.add_paragraph()
            r_poc_label = p_poc.add_run("驗證成果：")
            apply_font(r_poc_label, size=FONT_SIZE_BODY)
            r_poc_content = p_poc.add_run(desc_poc)
            apply_font(r_poc_content, size=FONT_SIZE_BODY)

        # ── 圖解 ─────────────────────────────────────────────
        p_img_label = doc.add_paragraph()
        r_img_label = p_img_label.add_run("圖解：")
        apply_font(r_img_label, size=FONT_SIZE_BODY)

        inserted_any = False
        captions = item.get("image_captions", [])
        for idx_img, img_filename in enumerate(images):
            img_path = IMAGES_DIR / img_filename
            if img_path.exists():
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(str(img_path), width=Cm(13))
                    # 圖題
                    caption_text = captions[idx_img] if idx_img < len(captions) else img_filename
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_cap = p_cap.add_run(f"圖、{caption_text}")
                    apply_font(r_cap, size=FONT_SIZE_BODY)
                    inserted_any = True
                except Exception as e:
                    print(f"  ⚠ 無法嵌入圖片 {img_filename}: {e}")
            else:
                print(f"  ⚠ 圖片不存在: {img_path}")

        if not inserted_any:
            p_no_img = doc.add_paragraph()
            r_no_img = p_no_img.add_run("（規劃中，擬於實作階段補上功能示意圖）")
            apply_font(r_no_img, size=FONT_SIZE_BODY)

        # image_notes 暫不輸出（提案階段不需要詳細截圖說明）

        # ── 施作所需月數 ─────────────────────────────────────
        p_months = doc.add_paragraph()
        r_m_label = p_months.add_run("施作所需月數：")
        apply_font(r_m_label, size=FONT_SIZE_BODY)
        months_str = f"{months:.0f} 個月" if months == int(months) else f"{months} 個月"
        r_m_val = p_months.add_run(months_str)
        apply_font(r_m_val, size=FONT_SIZE_BODY)

        # ── 估價 ─────────────────────────────────────────────
        p_price = doc.add_paragraph()
        r_p_label = p_price.add_run("估價：")
        apply_font(r_p_label, size=FONT_SIZE_BODY)
        r_p_val = p_price.add_run(format_amount(total))
        apply_font(r_p_val, size=FONT_SIZE_BODY)


    # ── 五年期提示句（粗體）────────────────────────────────────
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        "※另請於上述項目中，挑選出可於五年內完成之項目，羅列於下表"
        "（如表格不敷使用，請自行增列）："
    )
    apply_font(r_note, size=FONT_SIZE_BODY, bold=True)

    add_para(doc)

    # ── 五年期規劃表格標題 ──────────────────────────────────────
    p_table_title = doc.add_paragraph()
    p_table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tt = p_table_title.add_run("五年期系統功能規劃")
    apply_font(r_tt, size=FONT_SIZE_BODY, bold=True)

    # ── 五年期規劃表格（欄寬比 1:3.2:2.1:2.1）─────────────────
    total_rows = 1 + max(len(five_year_items), 5)  # 表頭 + 至少5列資料
    table = doc.add_table(rows=total_rows, cols=4)
    set_table_borders(table)
    # 欄寬（共約 14.6 cm）
    col_widths = [1.74, 5.58, 3.65, 3.65]
    set_col_widths(table, col_widths)

    # 表頭
    headers = ["項次", "項目名稱", "施作所需月數", "估價"]
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(hdr)
        apply_font(r, size=FONT_SIZE_BODY, bold=True)

    # 資料列
    for row_i, fy_item in enumerate(five_year_items):
        row = table.rows[row_i + 1]
        vals = [
            str(row_i + 1),
            fy_item["name"],
            f"{fy_item['months']:.0f} 個月" if fy_item["months"] == int(fy_item["months"]) else f"{fy_item['months']} 個月",
            format_amount(fy_item["total"]),
        ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = cell.paragraphs[0].add_run(val)
            apply_font(r, size=FONT_SIZE_BODY)

    doc.save(str(OUTPUT_FILE))
    print(f"\n✓ 已產出：{OUTPUT_FILE}")
    print(f"  總項目數：{len(items)}")
    print(f"  五年內可完成：{len(five_year_items)} 項")
    total_all = sum(int(i["months"] * pricing[i["difficulty"]]) for i in items)
    print(f"  全案估價合計：新臺幣 {total_all:,.0f} 元")


if __name__ == "__main__":
    main()
